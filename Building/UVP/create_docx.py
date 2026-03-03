"""
Generate MI UVP Google Docs-compatible DOCX
Defne's 4-category clustering with ring visualization
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

# --- Style helpers ---
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_category_header(doc, title, color_rgb, subtitle=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"● {title}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = color_rgb
    if subtitle:
        run2 = p.add_run(f"  —  {subtitle}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(107, 114, 128)

def add_uvp(doc, number, title, desc, diff, uscore, ready_text, color_rgb, is_future=False):
    # Title line with U-Score dots
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)

    # Number + Title
    run_num = p.add_run(f"{number}. ")
    run_num.bold = True
    run_num.font.size = Pt(9)
    run_num.font.color.rgb = color_rgb

    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(9)
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    # Readiness + U-Score
    dots = "●" * uscore + "○" * (10 - uscore)
    run_meta = p.add_run(f"   [{ready_text}]  {dots} {uscore}/10")
    run_meta.font.size = Pt(7)
    run_meta.font.color.rgb = RGBColor(107, 114, 128)

    if is_future:
        run_dollar = p.add_run("  $")
        run_dollar.font.size = Pt(8)
        run_dollar.font.color.rgb = RGBColor(5, 150, 105)
        run_dollar.bold = True

    # Description
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(1)
    p2.paragraph_format.left_indent = Pt(14)
    run_desc = p2.add_run(desc)
    run_desc.font.size = Pt(8)
    run_desc.font.color.rgb = RGBColor(55, 65, 81)

    # Differentiator
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(3)
    p3.paragraph_format.left_indent = Pt(14)
    run_diff = p3.add_run(diff)
    run_diff.italic = True
    run_diff.font.size = Pt(7.5)
    run_diff.font.color.rgb = RGBColor(107, 114, 128)

# --- Colors ---
ORANGE = RGBColor(234, 88, 12)    # Ses Analizi
PINK = RGBColor(190, 24, 93)      # Auditory Character & Nörolojik
BLUE = RGBColor(37, 99, 235)      # Social Listening
RED = RGBColor(220, 38, 38)       # Müzik Önerisi
GRAY = RGBColor(107, 114, 128)    # Yol Haritası
MI = RGBColor(99, 102, 241)       # Accent

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
run = title.add_run("Musical Intelligence")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(15, 23, 42)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(2)
run = subtitle.add_run("Unique Value Proposition")
run.font.size = Pt(11)
run.font.color.rgb = MI

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.paragraph_format.space_after = Pt(4)
run = tagline.add_run("Müzik dinlerken beyninde ne olduğunu anlayan ve buna göre sana rehberlik eden dünyadaki tek platform.")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(55, 65, 81)

# Metrics line
metrics = doc.add_paragraph()
metrics.alignment = WD_ALIGN_PARAGRAPH.CENTER
metrics.paragraph_format.space_after = Pt(6)
run = metrics.add_run("448 bilimsel çalışma  ·  131 bilişsel ölçüm  ·  24 persona  ·  26 beyin bölgesi  ·  4 nörokimyasal  ·  0 doğrudan rakip")
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(107, 114, 128)

# Category legend
legend = doc.add_paragraph()
legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
legend.paragraph_format.space_after = Pt(8)

run = legend.add_run("● Ses Analizi  ")
run.font.size = Pt(8); run.font.color.rgb = ORANGE; run.bold = True
run = legend.add_run("● Auditory Character & Nörolojik  ")
run.font.size = Pt(8); run.font.color.rgb = PINK; run.bold = True
run = legend.add_run("● Social Listening  ")
run.font.size = Pt(8); run.font.color.rgb = BLUE; run.bold = True
run = legend.add_run("● Müzik Önerisi")
run.font.size = Pt(8); run.font.color.rgb = RED; run.bold = True

# Separator
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(0)
sep.paragraph_format.space_after = Pt(4)
run = sep.add_run("─" * 95)
run.font.size = Pt(6)
run.font.color.rgb = MI

# ============================================================
# CATEGORY 1: SES ANALİZİ (TURUNCU)
# ============================================================
add_category_header(doc, "SES ANALİZİ", ORANGE, "Müziği duyma, ayrıştırma ve ölçme")

add_uvp(doc, 1, "Müzikal DNA Keşfi",
    "Spotify hesabını bağla, 30 saniyede müzikal kişiliğini öğren. Dinleme geçmişin 5 müzikal gene dönüşür: keşif iştahı, detay hassasiyeti, dramatik eğilim, duygusal derinlik, zevk esnekliği.",
    "Spotify \"beğenilen şarkılar\" gösterir. MI neden beğendiğini söyler.",
    8, "Hazır", ORANGE)

add_uvp(doc, 2, "24 Müzikal Persona",
    "\"Simyacı\", \"Mimar\", \"Kaşif\", \"Çıpa\", \"Kinetik\" — 5 aileden 24 karakter tipi. Gen profiline göre hangi persona olduğunu keşfet.",
    "16Personalities gibi ama müzik için — ve gerçek dinleme verisine dayalı.",
    7, "Hazır", ORANGE)

add_uvp(doc, 3, "Şarkı Röntgeni",
    "Herhangi bir şarkıyı yükle — beyninin onu nasıl işlediğini gör. Saniyede 172 kez taranarak 131 farklı bilişsel durum hesaplanır: dikkat, hafıza, duygu, hareket dürtüsü, ödül, sürpriz.",
    "Shazam şarkıyı tanır. MI şarkının seni nasıl etkilediğini gösterir.",
    10, "Hazır", ORANGE)

add_uvp(doc, 5, "Derinlik Seçici",
    "Aynı analiz 3 katmanda: 6 boyut (herkes anlar) → 12 boyut (meraklılar için) → 24 boyut (uzmanlar için: beyin bölgesi, nörokimyasal).",
    "Freemium modele doğal eşleme: ücretsiz=6D, premium=12D, profesyonel=24D.",
    8, "Hazır", ORANGE)

add_uvp(doc, 6, "Anında Profil",
    "Dosya yükleme yok, kayıt formu yok. Spotify'a giriş yap — sistem top parçalarını tarayarak anında MI profili oluşturur. Sıfır sürtünme.",
    "Kullanıcı ilk 30 saniyede değer görür — \"aha anı\" hemen gerçekleşir.",
    7, "Hazır", ORANGE)

add_uvp(doc, 7, "Şarkı Karşılaştırma",
    "İki şarkıyı yan yana koy — 131 boyutta farklarını gör. \"Neden A şarkısı beni dans ettirirken B şarkısı hüzünlendiriyor?\"",
    "Hiçbir platformda böyle bir karşılaştırma aracı yok.",
    7, "Hazır", ORANGE)

add_uvp(doc, 8, "Bilimsel Güvenilirlik",
    "2.847 makale tarandı, 448 dahil edildi. 1.116 ampirik iddia, 634 etki büyüklüğü. Nature Neuroscience hedefli manuscript hazır.",
    "Brain.fm \"bilime dayalı\" der ama açıklamaz. MI 448 çalışmayı tek tek gösterebilir.",
    9, "Hazır", ORANGE)

add_uvp(doc, 15, "Tam Mobil Uygulama",
    "React + Three.js WebGL, 21 sayfa, Spotify entegrasyonu, TR/EN. Persona keşfi, analiz labı, sohbet, görselleştirme — tek uygulamada.",
    "Teknoloji hazır, pazar farkı içerikte.",
    5, "Yakın", ORANGE)

# ============================================================
# CATEGORY 2: AUDITORY CHARACTER & NÖROLOJİK (PEMBE)
# ============================================================
add_category_header(doc, "AUDITORY CHARACTER & NÖROLOJİK INTERACTION", PINK, "Beynin müziğe tepkisini modelleme")

add_uvp(doc, 9, "Canlı Beyin Haritası",
    "Şarkı çalarken 26 beyin bölgesinin aktivasyonunu 3D olarak canlı izle. Dopamin altın, serotonin mavi, opioid yeşil olarak parlar.",
    "Dünyada müzik dinlerken beyin aktivasyonunu gerçek zamanlı görselleştiren başka uygulama yok.",
    9, "Yakın", PINK)

add_uvp(doc, 10, "Duygu Zaman Çizelgesi",
    "Şarkının her saniyesinde 14 duygusal boyut: valans, uyarılma, gerilim, ürperme, nostalji, huşu, huzur... Timeline üzerinde görselleştirilir.",
    "Spotify \"enerji: yüksek\" der. MI şarkının duygusal hikayesini anlatır.",
    8, "Yakın", PINK)

add_uvp(doc, 11, "Groove ve Hareket Analizi",
    "Ritim sürüklenmesi, hareket dürtüsü, senkronizasyon gücü, groove yoğunluğu. \"Bu şarkının groove puanı 8.7 — baş sallama dürtüsü 2:10'da zirve.\"",
    "Spotify \"dans edilebilirlik\" verir. MI motor korteksin gerçekte ne yaptığını ölçer.",
    7, "Yakın", PINK)

add_uvp(doc, 13, "Ürperme Tahmini",
    "Şarkının sende \"chills\" yaratma olasılığını ve hangi anlarda olabileceğini tahmin eder. Opioid + dopamin + sürpriz etkileşimine dayalı.",
    "Blood & Zatorre 2001 araştırmasına dayanan hesaplamalı ürperme modeli — dünyada ilk.",
    8, "Yakın", PINK)

add_uvp(doc, 14, "Hatıra Defteri",
    "\"Bu şarkı neden farklı hissettiriyor?\" — Aşinalık gücü, tanıma güveni, anlam yükleme analizi.",
    "Hiçbir platform \"neden bu şarkıyı tekrar dinliyorsun\" sorusuna bilimsel cevap vermiyor.",
    7, "Yakın", PINK)

add_uvp(doc, 16, "Sürpriz Dedektörü",
    "Şarkıyı senin gibi \"bekleyerek\" dinler ve beklentinin kırıldığı anları işaretler. \"2:42'deki modülasyon beyninin beklediği çözülmenin tersini yapıyor.\"",
    "Müzik eleştirmenlerinin sezgisel yaptığını ilk kez hesaplamalı yapan sistem.",
    9, "3-6 ay", PINK, is_future=True)

add_uvp(doc, 17, "Dopamin Haritası",
    "Beyninin hangi anlarda \"istiyorum\" (wanting) ve \"beğeniyorum\" (liking) sinyali gönderdiğini gör. 4 nörokimyasalın anlık seviyesi.",
    "Müzik dinlerken nörokimyasal dinamikleri modelleyen dünyada başka hiçbir ürün yok.",
    10, "3-6 ay", PINK, is_future=True)

add_uvp(doc, 19, "Dikkat Radarı",
    "Şarkıda dikkatinin nereye kaydığını gör. Çalışma müziği seçerken hangi şarkıların dikkatini dağıtmadan arka planda kaldığını bil.",
    "Brain.fm müzik üretir. MI herhangi bir şarkının dikkat etkisini ölçer.",
    8, "3-6 ay", PINK, is_future=True)

# ============================================================
# CATEGORY 3: SOCIAL LISTENING (MAVİ)
# ============================================================
add_category_header(doc, "SOCIAL LISTENING", BLUE, "Paylaşımlı dinleme ve sosyal müzik deneyimi")

add_uvp(doc, 4, "AI Müzik Danışmanı",
    "\"Bu şarkı neden beni ağlatıyor?\" diye sor, bilimsel cevap al. 11 araçla donatılmış kişisel yapay zeka. Personana göre konuşur.",
    "ChatGPT genel bilgi verir. MI ajanı şarkıyı senin beynin üzerinden analiz eder.",
    9, "Hazır", BLUE)

add_uvp(doc, 12, "Duo: Birlikte Dinle",
    "Arkadaşınla aynı şarkıyı dinle — müzikal uyumu gör. Combo sistemi, başarımlar, görevlerle oyunlaştırılmış.",
    "Spotify Blend playlist yapar. MI iki beynin müziğe nasıl farklı tepki verdiğini gösterir.",
    8, "Yakın", BLUE)

add_uvp(doc, 21, "Müzikal Evrim Takibi",
    "Haftalık/aylık dinleme verisinden zevkinin nasıl değiştiğini izle. \"Son 3 ayda keşif iştahın %18 arttı.\"",
    "Müzikal zevk değişimini bilimsel olarak takip eden hiçbir ürün yok.",
    9, "6-12 ay", BLUE, is_future=True)

add_uvp(doc, 22, "Sosyal Rezonans",
    "Konserde, partide, arabada — grup dinleme deneyimini ölç. \"Salonun %73'ü aynı anda ürperme yaşadı.\" Arkadaş grubu için müzikal uyumluluk skoru.",
    "Müzik dinlemenin sosyal boyutunu modelleyen dünyadaki tek girişim.",
    8, "6-12 ay", BLUE, is_future=True)

# ============================================================
# CATEGORY 4: MÜZİK ÖNERİSİ (KIRMIZI)
# ============================================================
add_category_header(doc, "MÜZİK ÖNERİSİ", RED, "Kişisel beyin profiline dayalı yönlendirme")

add_uvp(doc, 18, "Ruh Hali Navigatörü",
    "\"Stresli hissediyorum, sakinleşmek istiyorum\" → senin kognitif profiline göre geçişi sağlayacak şarkı sırası önerir.",
    "Spotify playlist'i herkese aynı. MI'nın önerisi kognitif profile özel.",
    9, "3-6 ay", RED, is_future=True)

add_uvp(doc, 20, "Çalışma Müziği Optimizer",
    "Kütüphanenden dikkat dağıtmayan, odak artıran şarkıları otomatik filtrele. \"Senin beynin için en iyi çalışma şarkıları\" listesi.",
    "Kişiye özel odak müziği — üretilmiş değil, senin kütüphanenden seçilmiş.",
    8, "3-6 ay", RED, is_future=True)

add_uvp(doc, 23, "Uyku/Rahatlama Rehberi",
    "Hangi şarkılar serotonin yükseltiyor, hangileri uyarılma düşürüyor? Kişisel \"uyku müziği\" listesi — üretilmiş değil, gerçek müzikten seçilmiş.",
    "Endel yapay ses üretiyor. MI senin müziğinden en uygununu seçiyor.",
    7, "6-12 ay", RED, is_future=True)

# ============================================================
# YOL HARİTASI
# ============================================================
add_category_header(doc, "YOL HARİTASI", GRAY, "Tasarım aşamasında — gelecek özellikler")

add_uvp(doc, 24, "Terapi Modu",
    "Anksiyete, stres, depresyon, uyku — kişinin kognitif profiline göre kalibre edilmiş müzik seçimi. Klinik doğrulama ile.",
    "LUCID ve Rubato Life sensör gerektiriyor. MI sadece ses ile çalışır.",
    10, "Yol Haritası", GRAY, is_future=True)

add_uvp(doc, 25, "Müzik Üretimi (Compose)",
    "\"Beni şu hissettirsin\" de — MI o müziği üretsin. Hedef kognitif durum → müzik.",
    "Suno/Udio metinden müzik yapar. MI, hedef beyin durumundan müzik yapar.",
    9, "Yol Haritası", GRAY, is_future=True)

add_uvp(doc, 26, "Eş-Yaratım (Hybrid)",
    "Müzik yaparken MI sana gerçek zamanlı geri bildirim verir. \"Bu nota gerilimi %23 artırdı, dopamin beklentisi oluşturdu.\"",
    "Hiçbir DAW kognitif geri bildirim vermiyor.",
    10, "Yol Haritası", GRAY, is_future=True)

add_uvp(doc, 27, "Konser/Canlı Analiz",
    "Canlı performansı gerçek zamanlı analiz et. \"Seyirci 2. şarkıda dikkat kaybetti, 5. şarkıda doruk noktası.\" Set listesi optimizasyonu.",
    "Konser deneyimini kognitif olarak analiz eden ilk araç.",
    8, "Yol Haritası", GRAY, is_future=True)

add_uvp(doc, 28, "Eğitim Modu",
    "Müzik öğrencileri için: \"Bu parçayı çalarken dinleyicinin dikkati 3. ölçüde düşüyor — dinamik kontrastla düzelt.\"",
    "Müzik eğitiminde dinleyici perspektifinden geri bildirim veren ilk sistem.",
    7, "Yol Haritası", GRAY, is_future=True)

# ============================================================
# REKABET
# ============================================================
sep2 = doc.add_paragraph()
sep2.paragraph_format.space_before = Pt(8)
sep2.paragraph_format.space_after = Pt(4)
run = sep2.add_run("─" * 95)
run.font.size = Pt(6)
run.font.color.rgb = MI

comp_title = doc.add_paragraph()
comp_title.paragraph_format.space_after = Pt(4)
run = comp_title.add_run("NEDEN BAŞKA KİMSE YAPAMIYOR?")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(15, 23, 42)

competitors = [
    ("Suno, Udio, Lyria 3", "Müzik üretiyor. MI anlıyor."),
    ("Spotify, Apple Music", "Şarkı öneriyor. MI neden sevdiğini açıklıyor."),
    ("Bridge.audio, Cyanite", "Etiket yapıştırıyor. MI beyin haritası çıkarıyor."),
    ("Brain.fm, Endel", "Ses üretiyor. MI her müziği analiz ediyor."),
    ("LUCID, Rubato Life", "Sensör gerektiriyor. MI sadece ses ile."),
    ("C-BMI (Japonya)", "EEG gerektiriyor. MI telefondan çalışıyor."),
]

table = doc.add_table(rows=len(competitors)+1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, text in enumerate(["Rakipler", "MI Farkı"]):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, "0F172A")

for row_idx, (comp, diff) in enumerate(competitors, 1):
    for col_idx, text in enumerate([comp, diff]):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(55, 65, 81)
        if row_idx % 2 == 0:
            set_cell_shading(cell, "F8FAFC")

# Closing
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing.paragraph_format.space_before = Pt(10)
run = closing.add_run("448 çalışma. 131 ölçüm. 96 nörobilim modeli. 756 dosya kod.")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(15, 23, 42)

closing2 = doc.add_paragraph()
closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing2.paragraph_format.space_after = Pt(2)
run = closing2.add_run("0 doğrudan rakip.")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = MI

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Mart 2026  ·  Şirket İçi  ·  Musical Intelligence Platform")
run.font.size = Pt(7)
run.font.color.rgb = RGBColor(156, 163, 175)

# Save
output = "/Volumes/SRC-9/SRC Musical Intelligence/Building/UVP/MI-UVP-2026-Defne.docx"
doc.save(output)
print(f"Saved to {output}")
